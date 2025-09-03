import os
import re
import psycopg2
from flask import Flask, request, jsonify
import pdfplumber
from docx import Document


from config import DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD, DATA_DIR, get_embedding
from model import answer_question
from filling_database import filling_database

app = Flask(__name__, static_folder='../app/build', static_url_path='/')


conn = psycopg2.connect(
    host=DB_HOST,
    port=DB_PORT,
    dbname=DB_NAME,
    user=DB_USER,
    password=DB_PASSWORD
)
cursor = conn.cursor()

def get_table_name(filename):
    base = os.path.splitext(filename)[0]
    return f"items_{base.replace('-', '_').replace(' ', '_')}"

def is_table_empty(table_name):
    cursor.execute(f"""
        SELECT EXISTS (
            SELECT FROM information_schema.tables 
            WHERE table_name = %s
        )
    """, (table_name,))
    table_exists = cursor.fetchone()[0]
    if not table_exists:
        print(f"Table '{table_name}' does not exist.")
        print('Creating table...')
        cursor.execute("""
            create extension if not exists vector;
        """)
        cursor.execute(f"""
            create table {table_name} (
                id serial primary key,
                content text,
                embedding vector(4096)
            );
        """)
        conn.commit()
        return True
    cursor.execute(f'SELECT COUNT(*) FROM {table_name}')
    count = cursor.fetchone()[0]
    return count == 0

@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    question = data.get('question', '')
    table = data.get('table', None)
    cursor.execute("""
        SELECT table_name FROM information_schema.tables WHERE table_name LIKE 'items_%'
    """)
    all_tables = [row[0] for row in cursor.fetchall()]
    found_table = None
    q_lower = question.lower()
    print(f"DEBUG: all_tables={all_tables}")
    for t in all_tables:
        fname_raw = t.replace('items_', '')
        fname_variants = set()
        fname_variants.add(fname_raw.lower())
        fname_variants.add(fname_raw.replace('_', ' ').lower())
        fname_variants.add(fname_raw.replace('_', '').lower())
        if '.' in fname_raw:
            fname_noext = fname_raw.split('.')[0]
            fname_variants.add(fname_noext.lower())
            fname_variants.add(fname_noext.replace('_', ' ').lower())
            fname_variants.add(fname_noext.replace('_', '').lower())
        for variant in fname_variants:
            if variant in q_lower:
                found_table = t
                print(f"DEBUG: matched table {t} with variant '{variant}' in question '{question}'")
                break
        if found_table:
            break
    if found_table:
        table = found_table
    if not question:
        return jsonify({'error': 'No question provided'}), 400
    if not table:
        return jsonify({'error': 'No table specified and no filename found in question'}), 400
    from model import answer_question
    import io
    import sys
    old_stdout = sys.stdout
    sys.stdout = mystdout = io.StringIO()
    answer_question(question, table)
    sys.stdout = old_stdout
    output = mystdout.getvalue()
    answer = output.split('The answer:')[-1].strip() if 'The answer:' in output else output.strip()
    return jsonify({'question': question, 'answer': answer})

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    base_filename = os.path.splitext(file.filename)[0]
    save_path = os.path.join(DATA_DIR, base_filename + '.md')

    if ext == '.pdf':
        with pdfplumber.open(file) as pdf:
            text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(text)
    elif ext in ['.docx', '.doc']:
        doc = Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(text)
    elif ext == '.md' or ext == '.txt':
        file.save(os.path.join(DATA_DIR, base_filename + ext))
    else:
        return jsonify({'error': 'Unsupported file type'}), 400

    table_name = get_table_name(file.filename)
    is_table_empty(table_name) 
    if ext == '.md' or ext == '.txt':
        saved_filename = base_filename + ext
    else:
        saved_filename = base_filename + '.md'
    filling_database(DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD, DATA_DIR, get_embedding, table_name, saved_filename)

    return jsonify({'message': 'File uploaded and database updated.', 'filename': base_filename, 'table': table_name})

@app.route('/status', methods=['GET'])
def status():
    files = [f for f in os.listdir(DATA_DIR) if os.path.isfile(os.path.join(DATA_DIR, f))]
    tables = []
    for f in files:
        tables.append(get_table_name(f))
    return jsonify({'tables': tables, 'files': files})


@app.route('/delete_file', methods=['POST'])
def delete_file():
    data = request.get_json()
    filename = data.get('filename')
    if not filename:
        return jsonify({'error': 'No filename provided'}), 400
    file_path = os.path.join(DATA_DIR, filename)
    file_deleted = False
    table_deleted = False
    error_msg = None
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            file_deleted = True
        except Exception as e:
            error_msg = f'Could not delete file: {str(e)}'
    else:
        error_msg = 'File not found.'
    # Видалення таблиці
    table_name = get_table_name(filename)
    cursor.execute("""
        SELECT EXISTS (
            SELECT FROM information_schema.tables WHERE table_name = %s
        )
    """, (table_name,))
    table_exists = cursor.fetchone()[0]
    if table_exists:
        try:
            cursor.execute(f"DROP TABLE IF EXISTS {table_name}")
            conn.commit()
            table_deleted = True
        except Exception as e:
            error_msg = f'Could not delete table: {str(e)}'
    else:
        # Якщо таблиці немає, це не помилка
        table_deleted = True
    # Формуємо відповідь
    if file_deleted and table_deleted:
        return jsonify({'message': f'File {filename} and table {table_name} deleted.'})
    elif file_deleted and not table_deleted:
        return jsonify({'error': error_msg or 'Table not deleted.'}), 500
    elif not file_deleted and table_deleted:
        return jsonify({'error': error_msg or 'File not deleted.'}), 500
    else:
        return jsonify({'error': error_msg or 'File and table not deleted.'}), 500

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def index(path):
    return app.send_static_file('index.html')

if __name__ == '__main__':
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    app.run(host='localhost', port=5000)
